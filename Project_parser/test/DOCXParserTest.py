#! /usr/bin/env python
# -*- coding: utf-8 -*-

import unittest
from Project_parser.parsers import DOCXParser
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx import Document
from docx.shared import Inches


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


class DocParserGetTextTest(unittest.TestCase):

    def test_empty_file(self):
        filename = '../.data/empty.docx'
        document = Document()
        document.save(filename)
        parser = DOCXParser.DOCXParser(filename)
        text = parser.get_text()
        self.assertEqual("", text)

    def test_simple_text(self):
        filename = '../.data/hello_world.docx'
        document = Document()
        document.add_paragraph('Hello, world!')
        document.save(filename)
        parser = DOCXParser.DOCXParser(filename)
        text = parser.get_text()
        self.assertEqual('Hello, world!', text)

    def test_ascii_text(self):
        filename = '../.data/ascii.docx'
        parser = DOCXParser.DOCXParser(filename)
        text = parser.get_text()
        self.assertEqual('╚◙Ї§○ї', text)

    def test_russian(self):
        filename = '../.data/russian.docx'
        document = Document()
        document.add_paragraph('Привет, мир!')
        document.save(filename)
        parser = DOCXParser.DOCXParser(filename)
        text = parser.get_text()
        self.assertEqual('Привет, мир!', text)

    def test_multiple_lines(self):
        filename = '../.data/multiple_lines.docx'
        document = Document()
        document.add_heading('Document Title', 0)
        p = document.add_paragraph('Example ')
        p.add_run('bold').bold = True
        p.add_run(' and some ')
        p.add_run('italic.').italic = True
        document.add_heading('Heading, level 1', level=1)
        document.add_paragraph('Intense quote', style='Intense Quote')
        document.save(filename)
        parser = DOCXParser.DOCXParser(filename)
        text = parser.get_text()
        self.assertEqual('Document Title\n\nExample bold and some italic.\n\nHeading, level 1\n\nIntense quote', text)

    def test_special_symbols(self):
        filename = '../.data/special_symbols.docx'
        test_text = "`@#$%^&*()_+-=[]{}:;\"|\\'<,>.?/"
        document = Document()
        document.add_paragraph(test_text)
        document.save('filename')
        parser = DOCXParser.DOCXParser("filename")
        text = parser.get_text()
        self.assertEqual(test_text, text)

    def test_with_image(self):
        filename = '../.data/with_image.docx'
        document = Document()
        document.add_paragraph('Hello, world!')
        document.add_picture('.data/korablik-parusnik-model.jpg', width=Inches(1.25))
        document.add_paragraph('Hello!')
        document.save(filename)
        parser = DOCXParser.DOCXParser(filename)
        text = parser.get_text()
        self.assertEqual('Hello, world!\n\n\n\nHello!', text)

    def test_link(self):
        filename = '../.data/link.docx'
        document = Document()
        document.add_paragraph('Hello, world!')
        p = document.add_paragraph('Hello! ')
        add_hyperlink(p, 'Link', "spbu.ru")
        document.save(filename)
        parser = DOCXParser.DOCXParser(filename)
        text = parser.get_text()
        link = parser.get_links()
        self.assertEqual('Hello, world!\n\nHello! Link', text)
        self.assertEqual(['spbu.ru'], link)

    def test_multiple_links(self):
        filename = '../.data/multiple_links.docx'
        document = Document()
        document.add_paragraph('Hello, world!')
        p = document.add_paragraph('Hello! ')
        add_hyperlink(p, 'Link1', "spbu.ru")
        p = document.add_paragraph('\n\nParagraph with another link ')
        add_hyperlink(p, 'Link2', "http://www.apmath.spbu.ru")
        document.save(filename)
        parser = DOCXParser.DOCXParser(filename)
        link = parser.get_links()
        text = parser.get_text()
        self.assertEqual(["spbu.ru", "http://www.apmath.spbu.ru"], link)
        self.assertEqual('Hello, world!\n\nHello! Link1\n\n\n\nParagraph with another link Link2', text)

    def test_links_in_text(self):
        filename = '../.data/links_in_text.docx'
        document = Document()
        document.add_paragraph('Hello, world!')
        p = document.add_paragraph('Hello! ')
        add_hyperlink(p, 'Link1', "spbu.ru")
        document.add_paragraph('Paragraph with link (http://www.apmath.spbu.ru) in text')
        document.save(filename)
        parser = DOCXParser.DOCXParser(filename)
        link = parser.get_links()
        text = parser.get_text()
        self.assertEqual(["spbu.ru"], link)
        self.assertEqual('Hello, world!\n\nHello! Link1\n\nParagraph with link (http://www.apmath.spbu.ru) in text', text)


if __name__ == '__main__':
    unittest.main()
